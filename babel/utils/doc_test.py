from django.shortcuts import get_object_or_404
import os
from docx import Document
from docx.oxml.ns import qn

from appraisal.models import AppraisalFileRecord, BasicInfo, \
    AppraisalInfo, FilePhase, AppraisalFile

from babel.settings import BASE_DOC_FOLDER, DEST_DOC_FOLDER

"""
生成报告的流程：
1 - 从数据库获取所有字段的数据，保存在一个dict中。key是模板docx文件中待替换的编号，value是要替换的文本。
2 - 每一个要生成的文件，都对应一个独立的方法，这样便于排错。
3 - 每次读取一个文件，替换里面的内容，然后保存到新的文件。
4 - 所有文件生成后，打包成一个zip文件，路径记录在数据库中，和项目对应。
"""

"""
数据说明：
A02: 鉴定类别                 文件:(2/18/)
A03: 项目编号                 文件:(2/4/6/7/11/12/13/14/15/16/18/19)
A04: 案由                    文件:(2/4/7/16/)
A05: 委托人                   文件:(2/4/6/7/11/14/15/)              
A06: 委托事项                 文件:(2/4/5/6/7/8/11/12/14/17/18/)
A07: 是否重新鉴定              文件:(2/4/5/)
A08: 鉴定对象                  文件:(2/7/9/11/15/17/)
A09: 受理时间                  文件:(2/4/5/6/7/11/14/)
A10: 完成时间                  文件:(2/7/11/12/17/19)
A11: 鉴定人(多个)              文件:(2/10/13/17/18/)
A12: 复核人                    文件:(2/)
A13: 主要鉴定意见               文件:(2/10/12/13/)
A14: 立卷人                    文件:(2/19)
A15: 归档日期                  文件:(2/)
A16: 鉴定地址                  文件:(5/8/9/11/17/)
A17: 基本案情                   文件:(4/10/)
A18: 鉴定材料(多个)              文件:(4/6/7/11/14/15/)
A19: 送达情况                   文件:(4/16/)
A21: 联系人                     文件:(4/)
A22: 电话                       文件:(4/)
"""


class ApprFile:

    def __init__(self, filename, quantity, received_date, receiver, records):
        self.filename = filename
        self.quantity = quantity
        self.received_date = received_date
        self.receiver = receiver
        self.records = records

    def __str__(self):
        return "文件名：" + str(self.filename) \
               + "\n数量：" + str(self.quantity) \
               + "\n接收日期：" + str(self.received_date) \
               + "\n接收人：" + str(self.receiver)


def data_preparer(basic_info_id):
    """
    准备数据，读取数据库，将指定项目的所有数据存入字典中。
    所有value全部为字符串，或者字符串组成的list。
    :param basic_info_id: 项目ID
    :return: 数据组成的字典。
    """

    data = {}
    # 获取数据
    basic_info = get_object_or_404(BasicInfo, id=basic_info_id)
    file_phase = get_object_or_404(FilePhase, basic_info=basic_info)
    appraisal_info = get_object_or_404(AppraisalInfo, basic_info=basic_info)

    # 单独处理鉴定材料，鉴定材料包含多个属性，封装成对象，添加到列表中。
    appraisal_file = AppraisalFile.objects.filter(basic_info=basic_info)
    appraisal_file_list = []
    appr_file_name_list = []
    for item in appraisal_file:
        records = AppraisalFileRecord.objects.filter(appraisal_file=item)
        appraisal_file_list.append(ApprFile(item.name, item.quantity, item.received_date,
                                            item.receiver.name, records))
        appr_file_name_list = [item.filename for item in appraisal_file_list]

    data["appraisal_file_list"] = appraisal_file_list
    data["appr_file_name_list"] = appr_file_name_list

    data["A18"] = "\n".join(appr_file_name_list)

    data["A02"] = basic_info.type.name
    # print("A02: ", data["A02"])

    data["A03"] = basic_info.sn
    # print("A03: ", data["A03"])

    data["A04"] = basic_info.purpose.name
    # print("A04: ", data["A04"])

    data["A05"] = basic_info.principal
    # print("A05: ", data["A05"])

    data["A06"] = basic_info.trust_detail
    # print("A06: ", data["A06"])

    data["A07"] = "是" if basic_info.is_re_appraisal else "否"
    # print("A07: ", data["A07"])

    data["A08"] = basic_info.target
    # print("A08: ", data["A08"])

    data["A09"] = str(basic_info.created_date)
    # print("A09: ", data["A09"])

    data["A10"] = str(file_phase.finished_date)
    # print("A10: ", data["A10"])

    team = appraisal_info.appraisal_team.all()
    name_list = []
    for item in team:
        name_list.append(item.name)
    data["A11"] = "、".join(name_list)
    # print("A11: ", data["A11"])

    data["A12"] = appraisal_info.final_reviewer.name
    # print("A12: ", data["A12"])

    data["A13"] = appraisal_info.opinion
    # print("A13: ", data["A13"])

    data["A14"] = appraisal_info.archivist.name
    # print("A14: ", data["A14"])

    data["A15"] = str(file_phase.file_date)
    # print("A15: ", data["A15"])

    data["A16"] = appraisal_info.appraisal_address
    # print("A16: ", data["A16"])

    data["A17"] = appraisal_info.project_detail
    # print("A17: ", data["A17"])

    # appraisal_file_list = []
    # for item in appraisal_file:
    #     appraisal_file_list.append(item.name)
    # data["A18"] = "、".join(appraisal_file_list)
    # print("A18: ", data["A18"])

    data["A19"] = file_phase.delivery.name
    # print("A19: ", data["A19"])

    data["A21"] = appraisal_info.contact
    # print("A21: ", data["A21"])

    data["A22"] = appraisal_info.phone
    # print("A22: ", data["A22"])

    return data


def file_maker(data_dict):
    dest_dir = DEST_DOC_FOLDER
    for root, dirs, files in os.walk(BASE_DOC_FOLDER):
        for file in files:
            # 排除special开头的文件，目前只有15号文件，表格有多行，需要完全重新生成。
            if not file.startswith("special"):
                source_file = os.path.join(root, file)
                dest_file = os.path.join(dest_dir, file)
                document = Document(source_file)

                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for key in data_dict:
                                if key in cell.text and not isinstance(data_dict[key], list):
                                    cell.text = cell.text.replace(key, data_dict.get(key))
                for para in document.paragraphs:
                    for run in para.runs:
                        for key in data_dict:
                            if key in run.text and not isinstance(data_dict[key], list):
                                run.text = run.text.replace(key, data_dict[key])
                document.save(dest_file)


def file_15_maker(file_15_data_dict):
    dest_dir = DEST_DOC_FOLDER
    dest_file = os.path.join(dest_dir, "special-15-司法鉴定委托材料收领流转表.docx")
    document = Document()

    # 设置字体
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 添加标题
    document.add_heading("山东求是建筑工程司法鉴定所")
    document.add_heading("司法鉴定委托材料流转表", level=2)
    sn_para = document.add_paragraph("编号：" + file_15_data_dict["A03"])

    # 计算表格的行数，默认最少11行
    # appraisal_file_list
    # appr_file_name_list
    lines = 11
    for item in file_15_data_dict["appraisal_file_list"]:
        lines += 1
        for record in item.records:
            lines += 1
    print(lines)
    # 添加表格
    table = document.add_table(rows=lines, cols=4)
    # 添加边框
    table.style = 'Table Grid'
    # 写入第一行数据（委托人、被鉴定对象）
    table.cell(0, 0).text = "委托人"
    table.cell(0, 1).text = file_15_data_dict["A05"]
    table.cell(0, 2).text = "被鉴定对象"
    table.cell(0, 3).text = file_15_data_dict["A08"]
    # 添加第二行
    table.cell(1, 0).text = "送鉴材料"
    table.cell(1, 0).merge(table.cell(1, 1)).merge(table.cell(1, 2)).merge(table.cell(1, 3))
    # 添加第三行
    table.cell(2, 0).text = "材料名称"
    table.cell(2, 1).text = "数量"
    table.cell(2, 2).text = "接收时间"
    table.cell(2, 3).text = "接收人"
    # 添加材料列表
    # 遍历材料列表，起始行号从3开始写入
    line_num = 3
    for item in file_15_data_dict["appraisal_file_list"]:
        table.cell(line_num, 0).text = item.filename
        table.cell(line_num, 1).text = str(item.quantity)
        table.cell(line_num, 2).text = str(item.received_date)
        table.cell(line_num, 3).text = item.receiver
        line_num += 1

    # 写入第五行
    table.cell(line_num, 0).text = "鉴定材料在鉴定机构内部流转记录"
    table.cell(line_num, 0).merge(table.cell(line_num, 1)).merge(table.cell(line_num, 2)).merge(table.cell(line_num, 3))
    line_num += 1
    # 写入第六行
    table.cell(line_num, 0).text = "鉴定材料名称"
    table.cell(line_num, 0).merge(table.cell(line_num, 1))
    table.cell(line_num, 2).text = "时间"
    table.cell(line_num, 3).text = "（签名）"
    line_num += 1
    # 写入鉴定材料借阅记录
    # 每条借阅记录占用两行
    for item in file_15_data_dict["appraisal_file_list"]:
        for record in item.records:
            table.cell(line_num, 0).text = item.filename
            # 合并上下两行
            table.cell(line_num, 0).merge(table.cell(line_num + 1, 0)).merge(table.cell(line_num + 1, 0)).merge(
                table.cell(line_num + 1, 1))
            table.cell(line_num, 2).text = str(record.borrowing_time)
            table.cell(line_num, 3).text = record.borrower.name
            table.cell(line_num + 1, 2).text = str(record.return_time)
            table.cell(line_num + 1, 3).text = record.borrower.name
            line_num += 2

    table.cell(line_num, 0).text = "退还委托人送鉴材料"
    table.cell(line_num, 0).merge(table.cell(line_num, 1)).merge(table.cell(line_num, 2)).merge(table.cell(line_num, 3))
    line_num += 1
    table.cell(line_num, 0).text = "材料名称"
    table.cell(line_num, 1).text = "数量"
    table.cell(line_num, 2).text = "接收时间"
    table.cell(line_num, 3).text = "接收人(签名)"
    line_num += 1

    for item in file_15_data_dict["appraisal_file_list"]:
        table.cell(line_num, 0).text = item.filename
        table.cell(line_num, 1).text = str(item.quantity)
        line_num += 1
    # 保存文件
    document.save(dest_file)

    # 路径写入数据库


    
