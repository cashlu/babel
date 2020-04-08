import os

from docx import Document
import re
from appraisal.models import BasicInfo, AppraisalInfo, FilePhase, LocaleFile
from django.shortcuts import get_object_or_404


# def replace_content(document, old_str, new_str):
#     for para in document.paragraphs:
#         for run in para.runs:
#             run.text = run.text.replace(old_str, new_str)
#
#     for table in document.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 if cell.text == old_str:
#                     cell.text = cell.text.replace(old_str, new_str)
#                     print('替换一次')
#     document.save('../../base_docs/dest.docx')


# def replace_content(document, word_context):
#     for para in document.paragraphs:
#         for run in para.runs:
#             if run.text in word_context:
#                 run.text = run.text.replace(run.text, word_context[run.text])
#     for table in document.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 if cell.text in word_context:
#                     cell.text = cell.text.replace(cell.text, word_context[cell.text])
#
#     document.save('../../base_docs/dest.docx')


def dict_generator(basic_info_id):
    """
    替换docx的dict的生成器函数.
    :return: 包含字段名和字段值的dict.
    """
    basic_info = get_object_or_404(BasicInfo, id=basic_info_id)

    # BasicInfo fields
    project_name = basic_info.name  # A00
    sn = basic_info.sn  # A03
    org = basic_info.org.name  # A01
    appraisal_type = basic_info.type.name  # A02
    purpose = basic_info.purpose.name  # A04
    principal = basic_info.principal  # A05
    trust_detail = basic_info.trust_detail  # A06
    is_re_appraisal = '是' if basic_info.is_re_appraisal else '否'  # A07
    target = basic_info.target  # A08
    prj_created_date = str(basic_info.created_date)  # A09
    trust_date = str(basic_info.trust_date)  # A23

    # AppraisalInfo fields
    appraisal_info = get_object_or_404(AppraisalInfo, basic_info=basic_info)
    appraisal_team = [item.name for item in appraisal_info.appraisal_team.all()]  # A11
    reviewer = appraisal_info.reviewer.name  # A12
    opinion = appraisal_info.opinion  # A13
    archivist = appraisal_info.archivist.name  # A14
    appraisal_address = appraisal_info.appraisal_address  # A16
    project_detail = appraisal_info.project_detail  # A17
    contact = appraisal_info.contact  # A21
    phone = appraisal_info.phone  # A22
    appraisal_date = str(appraisal_info.appraisal_date)  # A24
    discuss_date = str(appraisal_info.discuss_date)  # A25
    locale_files = [item.name for item in basic_info.localefile_set.all()]  # A18

    # FilePhase fields
    file_phase = get_object_or_404(FilePhase, basic_info=basic_info)
    finished_date = str(file_phase.finished_date)  # A10
    file_date = str(file_phase.file_date)  # A15
    delivery = file_phase.get_delivery_display()  # A19
    amount = file_phase.amount  # A20

    fields_dict = {
        'A00': project_name,
        'A01': org,
        'A02': appraisal_type,
        'A03': sn,
        'A04': purpose,
        'A05': principal,
        'A06': trust_detail,
        'A07': is_re_appraisal,
        'A08': target,
        'A09': prj_created_date,
        'A10': finished_date,
        'A11': appraisal_team,
        'A12': reviewer,
        'A13': opinion,
        'A14': archivist,
        'A15': file_date,
        'A16': appraisal_address,
        'A17': project_detail,
        'A18': locale_files,
        'A19': delivery,
        'A20': amount,
        'A21': contact,
        'A22': phone,
        'A23': trust_date,
        'A24': appraisal_date,
        'A25': discuss_date,
    }

    return fields_dict


def docx_replace_regex(doc_obj, regex, replace):
    """
    替换docx文件中的特定内容
    :param doc_obj: Document对象
    :param regex: 匹配的正则表达式模式
    :param replace: 替换的具体内容
    :return:
    """
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


# def docx_replace_regex(doc_obj, reg_dict):
#     """
#     替换docx文件中的特定内容
#     :param doc_obj: Document对象
#     :param reg_dict: 匹配的正则表达式模式
#     :return:
#     """
#     for p in doc_obj.paragraphs:
#         for regex, replace in reg_dict.items():
#             if regex.search(p.text):
#                 inline = p.runs
#                 # Loop added to work with runs (strings with same style)
#                 for i in range(len(inline)):
#                     print('regex: ' + str(regex))
#                     print('text : ' + inline[i].text)
#                     if regex.search(inline[i].text):
#                         text = regex.sub(replace, inline[i].text)
#                         inline[i].text = text
#
#     for table in doc_obj.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 # TODO: check
#                 docx_replace_regex(cell, reg_dict)


def files_generator(reg_dict):
    source_dir = 'base_docs'
    dest_dir = 'dest_docs'
    source_file_list = []
    for path, dir_list, file_list in os.walk(source_dir):
        for file in file_list:
            filepath = os.path.join(path, file)
            source_file_list.append(filepath)
    for file in source_file_list:
        print('processing -- ' + file)
        doc = Document(file)
        for key, value in reg_dict.items():
            docx_replace_regex(doc, re.compile(key), str(value))
            doc.save(os.path.join(dest_dir, os.path.basename(file)))
